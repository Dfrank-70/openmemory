from __future__ import annotations

import io
import subprocess
from pathlib import Path
from typing import Any

import openpyxl
import pdfplumber
from docx import Document

from src.config.settings import get_logger


def read_excel(path: str | Path) -> dict[str, Any]:
    workbook = openpyxl.load_workbook(filename=path, data_only=True)
    sheets_text: list[str] = []
    tables: list[dict[str, Any]] = []
    metadata = {"sheet_names": workbook.sheetnames}
    for sheet in workbook.worksheets:
        rows: list[str] = []
        for row in sheet.iter_rows(values_only=True):
            cells = ["" if cell is None else str(cell) for cell in row]
            rows.append(" | ".join(cells).strip())
        sheet_text = f"# Sheet: {sheet.title}\n" + "\n".join(row for row in rows if row)
        sheets_text.append(sheet_text)
        tables.append({"sheet": sheet.title, "rows": rows})
    return {"text": "\n\n".join(sheets_text), "metadata": metadata, "tables": tables}


def read_docx(path: str | Path) -> dict[str, Any]:
    document = Document(path)
    paragraphs = [paragraph.text.strip() for paragraph in document.paragraphs if paragraph.text.strip()]
    table_texts: list[str] = []
    tables: list[dict[str, Any]] = []
    for index, table in enumerate(document.tables, start=1):
        rows: list[str] = []
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells]
            rows.append(" | ".join(cells))
        table_texts.append(f"# Table {index}\n" + "\n".join(rows))
        tables.append({"table": index, "rows": rows})
    combined = "\n\n".join(paragraphs + table_texts)
    return {"text": combined, "metadata": {"paragraphs": len(paragraphs)}, "tables": tables}


def read_legacy_doc(path: str | Path) -> dict[str, Any]:
    logger = get_logger("file_readers")
    try:
        result = subprocess.run(["antiword", str(path)], capture_output=True, text=True, check=True)
        return {"text": result.stdout, "metadata": {"reader": "antiword"}, "tables": []}
    except Exception as exc:
        logger.warning("legacy_doc_reader_failed path=%s error=%s", path, exc)
        return {"text": "", "metadata": {"reader": "antiword", "error": str(exc)}, "tables": []}


def read_pdf(path: str | Path) -> dict[str, Any]:
    logger = get_logger("file_readers")
    text_chunks: list[str] = []
    tables: list[dict[str, Any]] = []
    with pdfplumber.open(path) as pdf:
        for page_number, page in enumerate(pdf.pages, start=1):
            page_text = page.extract_text() or ""
            if page_text.strip():
                text_chunks.append(f"# Page {page_number}\n{page_text}")
            extracted_tables = page.extract_tables() or []
            for table in extracted_tables:
                tables.append({"page": page_number, "rows": table})
        if not text_chunks:
            try:
                import pytesseract
                from PIL import Image
            except Exception as exc:
                logger.warning("ocr_dependencies_missing path=%s error=%s", path, exc)
            else:
                for page_number, page in enumerate(pdf.pages, start=1):
                    image = page.to_image(resolution=150).original
                    buffer = io.BytesIO()
                    image.save(buffer, format="PNG")
                    buffer.seek(0)
                    ocr_text = pytesseract.image_to_string(Image.open(buffer))
                    if ocr_text.strip():
                        text_chunks.append(f"# OCR Page {page_number}\n{ocr_text}")
    return {"text": "\n\n".join(text_chunks), "metadata": {"pages": len(text_chunks)}, "tables": tables}


def read_file_to_text(path: str | Path) -> dict[str, Any]:
    target = Path(path)
    suffix = target.suffix.lower()
    if suffix in {".xlsx", ".xlsm", ".xltx", ".xltm"}:
        return read_excel(target)
    if suffix == ".xls":
        return read_excel(target)
    if suffix == ".docx":
        return read_docx(target)
    if suffix == ".doc":
        return read_legacy_doc(target)
    if suffix == ".pdf":
        return read_pdf(target)
    text = target.read_text(encoding="utf-8")
    return {"text": text, "metadata": {"reader": "plain_text"}, "tables": []}
